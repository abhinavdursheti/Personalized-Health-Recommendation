import docx

doc_path = "Architecture Document.docx"
try:
    with open("arch_template_structure.txt", "w", encoding="utf-8") as f:
        doc = docx.Document(doc_path)
        
        f.write("--- PARAGRAPHS ---\n")
        for para in doc.paragraphs:
            if para.text.strip() != "":
                f.write(f"[{para.style.name}] {para.text}\n")
                
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for row in table.rows:
                row_data = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                f.write(str(row_data) + "\n")
            f.write("-" * 40 + "\n")
            
except Exception as e:
    with open("arch_template_structure.txt", "w", encoding="utf-8") as f:
        f.write(f"Error reading docx: {e}")
