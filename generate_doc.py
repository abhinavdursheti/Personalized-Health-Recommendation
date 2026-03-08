import docx
from docx.shared import Pt
import json

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)

def generate_functional_document():
    doc = docx.Document()
    
    # Title
    doc.add_heading("Functional Document - Personalized Health Recommendation System", level=0)
    
    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    doc.add_paragraph("The Personalized Health Recommendation System project aims to revolutionize individual health management by providing AI-powered, personalized recommendations for diet, exercise, and sleep. This sprint focuses on the implementation of advanced ML models to deliver Behavior Recovery & Stability Prediction, Behavior-Cause Correlation Analysis, and Habit Sensitivity Insights.")
    
    # 2. Product Goal 
    add_heading(doc, "2. Product Goal", 1)
    doc.add_paragraph("The primary goal of this project is to create an intelligent platform that analyzes user health data to generate highly personalized and actionable health plans. By identifying root causes of health changes and predicting recovery times, the system aims to provide users with a comprehensive and personalized fitness and wellness experience.")
    
    # 3. Demography
    add_heading(doc, "3. Demography (Users, Location)", 1)
    p_users = doc.add_paragraph()
    p_users.add_run("Users\n").bold = True
    p_users.add_run("Target Users: Health enthusiasts, individuals seeking weight management, fitness beginners, and nutrition-conscious individuals.\n")
    p_users.add_run("User Characteristics: Diverse age groups, varying fitness levels, and different pre-existing medical conditions or dietary restrictions.")
    
    p_location = doc.add_paragraph()
    p_location.add_run("Location\n").bold = True
    p_location.add_run("Target Location: Worldwide, accessible via web browser.")
    
    # 4. Business Processes
    add_heading(doc, "4. Business Processes", 1)
    doc.add_paragraph("The key business processes include:")
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("User Profiling and Onboarding: ").bold = True
    bp.add_run("Process for users to create secure accounts and securely log their health profile, including weight, age, activity level, and health goals.")
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("Health Data Tracking: ").bold = True
    bp.add_run("Process for users to log daily health metrics like sleep hours, exercise minutes, and calories consumed.")
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("ML-Powered Recommendation Generation: ").bold = True
    bp.add_run("Backend process utilizing Random Forest, Decision Tree, and Linear Regression models to evaluate user inputs and provide real-time diet, sleep, and exercise plans.")
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("Advanced Analytics Analysis: ").bold = True
    bp.add_run("Process of taking tracked historical data to perform root-cause correlation analysis, habit sensitivity profiling, and setback recovery predictions.")

    # 5. Features
    add_heading(doc, "5. Features", 1)
    doc.add_paragraph("The implementation focuses on the following key features:")
    
    add_heading(doc, "Feature #1: Intelligent Recommendation Engines", 2)
    doc.add_paragraph("Personalized Diet Plans (BMR/TDEE tracking), Workout Schedules, and Sleep hygiene models.")
    doc.add_paragraph().add_run("User Story:").bold = True
    doc.add_paragraph("As a new user seeking to improve my health, I want to receive tailored daily calorie and exercise plans so that I don't have to guess what workouts or meals fit my body and goals.")
    
    add_heading(doc, "Feature #2: Advanced Data Analytics Dashboard", 2)
    doc.add_paragraph("Real-time metric visualization, progress tracking charts, and dynamic insights.")
    doc.add_paragraph().add_run("User Story:").bold = True
    doc.add_paragraph("As a data-driven health enthusiast, I want to view my daily metrics plotted on interactive charts so that I can easily visually track my progress over the last month.")

    add_heading(doc, "Feature #3: Behavior & Stability Prediction Engine", 2)
    doc.add_paragraph("Uses ML to predict days needed to recover from setbacks and calculates behavior stability scores.")
    doc.add_paragraph().add_run("User Story:").bold = True
    doc.add_paragraph("As someone recovering from an injury or a period of poor health, I want to see a specific timeline for my recovery so that I can set realistic expectations and stay motivated.")

    add_heading(doc, "Feature #4: Habit Sensitivity & Root-Cause Analyzer", 2)
    doc.add_paragraph("Identifies primary causes of health changes and tracks fragile vs. resilient habits.")
    doc.add_paragraph().add_run("User Story:").bold = True
    doc.add_paragraph("As an individual struggling with weight fluctuations, I want the system to analyze my tracking history and pinpoint the exact habits causing these changes so that I can fix the root of the problem.")

    add_heading(doc, "Feature #5: Daily Health Tracking", 2)
    doc.add_paragraph("Process for users to log daily health metrics like sleep hours, exercise minutes, and calories consumed.")
    doc.add_paragraph().add_run("User Story:").bold = True
    doc.add_paragraph("As an active user, I want a simple daily entry form to quickly input my hours slept and exercise duration so that my ML models have fresh, accurate data to analyze.")

    # 6. Authorization Matrix
    add_heading(doc, "6. Authorization Matrix", 1)
    doc.add_paragraph("Define the roles and their corresponding access levels:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Access Level'
    
    roles = [
        ("Guest User", "Access to landing page and basic system information."),
        ("Standard User", "Full access to personal health profiling, daily tracking, dashboard analytics, and recommendation generation."),
        ("Admin", "Full access to user management, system database administration, and application infrastructure settings.")
    ]
    
    for role, access in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = access

    # 7. Assumptions
    add_heading(doc, "7. Assumptions", 1)
    doc.add_paragraph(style='List Bullet').add_run("Users will proactively log their daily health metrics consistently (minimum 7-10 days to achieve accurate analytics).")
    doc.add_paragraph(style='List Bullet').add_run("The pre-trained Machine Learning models (Random Forest, Decision Tree, etc.) maintain steady performance before requiring retaining on actual dynamic user datasets.")
    doc.add_paragraph(style='List Bullet').add_run("System infrastructure handles user data securely and compliantly.")
    
    # Save the document
    output_path = "Project_Functional_Document_Updated.docx"
    doc.save(output_path)
    print(f"Document saved successfully as {output_path}")

if __name__ == "__main__":
    generate_functional_document()
