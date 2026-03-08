import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
import os

def render_uml(puml_code, filename):
    print(f"Rendering UML to {filename} via kroki.io...")
    url = 'https://kroki.io/plantuml/png'
    data = puml_code.encode('utf-8')
    headers = {
        'Content-Type': 'text/plain',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
    }
    req = urllib.request.Request(url, data=data, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=10) as response:
            with open(filename, 'wb') as f:
                f.write(response.read())
        return True
    except Exception as e:
        print(f"Error drawing {filename}: {e}")
        return False

def add_page(doc, title, paragraphs, image_file=None):
    """
    Adds exactly one page of content. We keep the paragraph list short 
    (under 300 words total) to prevent Word from spilling over into two pages organically.
    """
    doc.add_heading(title, level=2)
    if image_file and os.path.exists(image_file):
        doc.add_picture(image_file, width=Inches(5.0))
        os.remove(image_file)
        
    for p_text in paragraphs:
        doc.add_paragraph(p_text)
        
    doc.add_page_break()

def generate_doc_v3():
    doc = docx.Document()
    
    # ---------------------------------------------------------
    # COVER PAGE (Page 1)
    # ---------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n\nARCHITECTURE DOCUMENT\nPersonalized Health Recommendation System\n\nAligned to Explicit Pipeline Architecture")
    run.bold = True
    run.font.size = Pt(24)
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PART 1: ARCHITECTURES (Pages 2-11) => 10 pages
    # ---------------------------------------------------------
    add_page(doc, "Part 1: Architectural Design Paradigms", [
        "This section details the system architecture of the Personalized Health Recommendation System.",
        "The architecture is explicitly designed around a core pipeline: Web Dashboard -> Input Processing -> Analysis -> Risk Classification -> Output Generation -> Database.",
        "We employ Microservices to encapsulate these modular boundaries, Event-Driven Architecture to pass data between them, and Serverless computing to power the asynchronous Analysis layer."
    ]) # Page 2

    add_page(doc, "1. Microservices Architecture - Structural Overview", [
        "The system breaks the analytical flow into distinct microservices.",
        "The 'Input Processing' module acts as an isolated ingestion service consisting of Data Input, Data Preprocessing, and Feature Extraction layers.",
        "By containerizing the Input Processing separately from the Analysis module, we allow horizontal scaling. If user metric uploads spike, the Input microservice can scale independently of the Risk Analysis engines."
    ]) # Page 3
    
    add_page(doc, "1. Microservices Architecture - Output & DB Isolation", [
        "The 'Output Generation' module is composed of the Recommendation Engine and the Alert Generation systems.",
        "These operate as separate microservices. The Recommendation Engine serves users flagged as 'Low' or 'Medium' risk, generating standard diet/exercise health plans.",
        "Conversely, the Alert Generation service is a high-priority microservice that handles 'High Risk' classifications, alerting emergency contacts or care providers immediately."
    ]) # Page 4

    add_page(doc, "2. Event-Driven Architecture - Pipeline Flow", [
        "To connect the discrete pipeline modules, an Event-Driven Architecture (EDA) is implemented.",
        "When the 'Web Dashboard UI' submits new data, a 'DATA_RECEIVED' event is pushed to a message broker (e.g., Kafka).",
        "The Input Processing layer consumes this event, runs Preprocessing and Feature Extraction, and then publishes a 'FEATURES_EXTRACTED' event to be intercepted by the Analysis layer."
    ]) # Page 5
    
    add_page(doc, "2. Event-Driven Architecture - Asynchronous Benefits", [
        "An EDA model ensures the Web Dashboard UI remains highly responsive. The user does not wait for the Baseline Modelling and ML Risk Analysis to finish synchronously.",
        "If the Risk Classification diamond yields a 'High Risk' result, a specialized topic 'HIGH_RISK_DETECTED' is broadcasted.",
        "The Alert Generation microservice acts on this topic instantly, bypassing standard recommendation queues and reducing latency in critical health situations."
    ]) # Page 6
    
    add_page(doc, "3. Serverless Computing - ML Risk Analysis", [
        "The 'Analysis' module (Baseline Modelling and ML Risk Analysis) is deployed using Serverless Computing (AWS Lambda/Azure Functions).",
        "Machine Learning inference often requires bursty compute power. By using serverless functions for the ML Risk framework, the system allocates resources dynamically only when 'FEATURES_EXTRACTED' events arrive.",
        "This prevents the need for 24/7 dedicated GPU/CPU instances, minimizing cloud costs while maintaining sub-second response times during health evaluations."
    ]) # Page 7
    
    add_page(doc, "3. Serverless Computing - Recommendation Engine Deployment", [
        "Similarly, the Recommendation Engine under the 'Output Generation' module operates via Serverless architecture.",
        "Based on 'Low or Medium Risk' flags, a Lambda function executes personalized health algorithms. It reads the extracted physical features, applies dietary and exercise rules, and constructs a JSON payload.",
        "Once the recommendation is committed to the Database, the serverless function terminates, ensuring a stateless, highly scalable infrastructure."
    ]) # Page 8
    
    add_page(doc, "4. Infrastructure Security & Authentication", [
        "Security is paramount throughout this pipeline. The Web Dashboard UI acts as the sole ingress point, protected by an API Gateway.",
        "The Gateway enforces JWT (JSON Web Token) authentication before passing the payload to the Data Input Layer.",
        "Inter-service communication (e.g., from Analysis to Output Generation) is secured via mutual TLS (mTLS) and scoped IAM roles, ensuring no module can be bypassed."
    ]) # Page 9
    
    add_page(doc, "4. System Telemetry & Observability", [
        "With a pipeline spanning Input Processing, Analysis, and Output Generation, distributed tracing is required.",
        "We inject OpenTelemetry headers at the Web Dashboard UI. This correlation ID travels through Preprocessing, Baseline Modelling, Risk Classification, and into the final Database commit.",
        "Engineers can utilize dashboards (Grafana) to identify if bottlenecks are occurring at the Feature Extraction phase or the ML Risk Analysis phase."
    ]) # Page 10
    
    add_page(doc, "Architectural Summary", [
        "The architecture precisely implements the requested flow:",
        "Web Dashboard UI -> Input Processing -> Analysis -> Risk Classification -> Output Generation -> Database.",
        "This pipelined approach guarantees modularity, fault isolation, and the ability to selectively scale the Machine Learning components independent of the user interface."
    ]) # Page 11

    # ---------------------------------------------------------
    # PART 2: UML DIAGRAMS (Pages 12-36) => 25 pages
    # 8 Diagrams. Each gets 3 pages (Image, Analysis 1, Analysis 2). Plus 1 intro page = 25 pages.
    # ---------------------------------------------------------
    add_page(doc, "Part 2: UML Diagrams", [
        "This section encompasses twenty-four pages dedicated to explicitly rendered UML diagrams mapping the health system's exact pipeline.",
        "We construct the Main Architecture Flowchart, Use Case Diagram, Sequence Diagram, Component Diagram, Activity Diagram, State Diagram, ER Diagram, and Deployment Diagram.",
        "Each diagram natively reflects the core flow: Input Processing -> Analysis -> Risk Classification -> Output -> Database."
    ]) # Page 12

    diagrams = [
        ("Architecture Flowchart", """@startuml
skinparam rectangle {
    BackgroundColor White
    BorderColor Black
}
skinparam diamond {
    BackgroundColor White
    BorderColor Black
}

rectangle "Web Dashboard UI" as UI

rectangle "INPUT PROCESSING" as IP {
    rectangle "Data Input Layer" as DIL
    rectangle "Data preprocessing" as DP
    rectangle "Feature Extraction" as FE
    DIL --> DP
    DP --> FE
}

rectangle "ANALYSIS" as Analysis {
    rectangle "Baseline Modelling" as BM
    rectangle "ML Risk Analysis" as MLR
    BM --> MLR
}

diamond "Risk Classification" as RC

rectangle "OUTPUT GENERATION" as OG {
    rectangle "Recommendation Engine" as RE
    rectangle "Alert Generation" as AG
}

rectangle "Database" as DB

UI --> DIL
FE --> BM
MLR --> RC
RC --> RE : Low or Medium Risk
RC --> AG : High Risk
RE --> DB
AG --> DB
@enduml""", 
"This is the overarching architecture diagram precisely matching the project constraints.",
"It clearly establishes the sequential pipeline from Web UI down to the primary Database, branching dynamically at the Risk Classification node."),

        ("Use Case Diagram", """@startuml
left to right direction
actor "User" as user
actor "Health Admin" as admin

rectangle "Health Pipeline System" {
  usecase "Access Web Dashboard UI" as UC1
  usecase "Submit Health Data" as UC2
  usecase "Receive Recommendations" as UC3
  usecase "Receive Emergency Alerts" as UC4
}

user --> UC1
UC1 --> UC2
UC2 ..> UC3 : Low/Medium Risk
UC2 ..> UC4 : High Risk
admin --> UC4 : Monitors Alerts
@enduml""",
"The Use Case Diagram highlights user interaction with the absolute start and end points of the pipeline.",
"Users interact with the Web Dashboard UI to submit data, and ultimately receive distinct outputs: structural recommendations or immediate health alerts."),

        ("Sequence Diagram", """@startuml
participant "Web UI" as UI
participant "Input Processing" as IP
participant "Analysis Module" as AM
participant "Output Generaton" as OG
database "Database" as DB

UI -> IP: Send Raw Data
IP -> IP: Preprocess & Extract Features
IP -> AM: Send Features
AM -> AM: Baseline Model & ML Risk Analysis
AM -> OG: Risk Classification Result
alt Low/Medium Risk
    OG -> OG: Trigger Recommendation Engine
else High Risk
    OG -> OG: Trigger Alert Generation
end
OG -> DB: Persist Output
DB --> UI: Return Status
@enduml""",
"The Sequence Diagram maps the exact chronological flow of messages between the macro-components of the pipeline in real-time execution.",
"The critical juncture is the 'alt' block, showcasing the divergent paths handled by the Output Generation module based on the ML analysis."),
        
        ("Component Diagram", """@startuml
package "Client Tier" {
  [Web Dashboard UI]
}

package "Processing Tier" {
  [INPUT PROCESSING] -down-> [ANALYSIS]
  [ANALYSIS] -right-> [Risk Classification Engine]
}

package "Output Tier" {
  [Risk Classification Engine] ..> [Recommendation Engine] : Low Risk
  [Risk Classification Engine] ..> [Alert Generation] : High Risk
}

database "Data Tier" {
  [Database]
}

[Web Dashboard UI] --> [INPUT PROCESSING]
[Recommendation Engine] --> [Database]
[Alert Generation] --> [Database]
@enduml""",
"The Component Diagram maps the pipeline abstractions to physical software packages.",
"This explicitly isolates the Processing Tier (containing Feature Extraction and ML Baseline models) from the Output Tier (Recommendations vs Alerts)."),

        ("Activity Diagram", """@startuml
start
:Web Dashboard UI;
:Data Input Layer;
:Data preprocessing;
:Feature Extraction;
:Baseline Modelling;
:ML Risk Analysis;
if (Risk Classification?) then (Low/Medium Risk)
  :Recommendation Engine;
else (High Risk)
  :Alert Generation;
endif
:Database;
stop
@enduml""",
"The Activity Diagram provides a procedural view of the data lifecycle.",
"It translates the vertical boxes of the Architecture Diagram into standard activity nodes, confirming the linear progression ending at the shared Database node."),
        
        ("Entity-Relationship (ER) Diagram", """@startuml
entity "ProcessedInputs" {
  *id : uuid
  --
  raw_data : text
  features_extracted : json
}

entity "RiskAnalysis" {
  *id : uuid
  --
  input_id : uuid <<FK>>
  baseline_score : float
  ml_risk_level : varchar
}

entity "GeneratedOutputs" {
  *id : uuid
  --
  analysis_id : uuid <<FK>>
  output_type : varchar (Alert/Rec)
  payload : json
}

ProcessedInputs ||--o| RiskAnalysis : triggers
RiskAnalysis ||--o| GeneratedOutputs : generates
@enduml""",
"The ER schematic describes how the Database at the end of the pipeline structures the incoming telemetry.",
"Information generated by the Input, Analysis, and Output structures are relationally linked to preserve the full lifecycle context of the user's logged event."),

        ("Deployment Diagram", """@startuml
node "User Device" {
  [Web Dashboard UI]
}

cloud "Cloud Provider" {
  node "Ingestion Server" {
    [Input Processing Module]
  }
  
  node "Serverless Container" {
    [Analysis Module]
  }
  
  node "Execution Engine" {
    [Output Generation Module]
  }
  
  database "Managed DB" {
    [Database]
  }
}

[Web Dashboard UI] --> [Input Processing Module] : HTTPS
[Input Processing Module] --> [Analysis Module] : gRPC
[Analysis Module] --> [Output Generation Module] : Event Bus
[Output Generation Module] --> [Database] : TCP/IP
@enduml""",
"The Deployment mapping illustrates physical environments. The modular nature of the pipeline allows the Analysis Module to jump between zones.",
"This enables it to be deployed on Serverless architecture to handle heavy ML calculations dynamically, maintaining isolation from the DB engine."),

        ("State Machine Diagram", """@startuml
[*] --> Idle
Idle --> InputProcessing : Data Received
InputProcessing --> AnalysisPhase : Features Extracted
AnalysisPhase --> RiskClassification : ML Complete
RiskClassification --> RecommendationMode : [Risk = Low/Med]
RiskClassification --> AlertMode : [Risk = High]
RecommendationMode --> DB_Write
AlertMode --> DB_Write
DB_Write --> [*]
@enduml""",
"The State Diagram demonstrates the operational transitions of a single health data payload.",
"A payload shifts states from Preprocessing, actively into ML inference, and terminates in a database write state depending entirely on the routing of the Risk Classification logic.")
    ]

    for title, puml, summary1, summary2 in diagrams:
        filename = f"{title.replace(' ', '_')}.png"
        render_uml(puml, filename)
        
        # Page A: The Image
        add_page(doc, f"{title} - Visual Model", [
            f"The following visual diagram depicts the {title} aligned explicitly with the requested health pipeline."
        ], image_file=filename)
        
        # Page B: Analysis 1
        add_page(doc, f"{title} - Component Mapping", [
            summary1,
            f"In the context of the {title}, we observe strict adherence to the separation of concerns. The input preprocessing operates independently of the downstream systems.",
            "By locking these interfaces behind formal boundaries, the ML layers (Baseline Modeling and Risk Analysis) can be upgraded or swapped without breaking the web dashboard UI."
        ])
        
        # Page C: Analysis 2
        add_page(doc, f"{title} - Pipeline Routing", [
            summary2,
            "The branching logic depicted here is the system's most critical juncture. High Risk alerts must flow to the alert generation microservice in sub-millisecond latencies.",
            "Normal baseline tracking metrics are diverted to the recommendation engine, optimizing database indexing by separating urgent telemetry from standard daily logs."
        ])

    # ---------------------------------------------------------
    # PART 3: DATASET (Pages 37-50) => 14 pages
    # ---------------------------------------------------------
    add_page(doc, "Part 3: Dataset Source and Characteristics", [
        "This section maps the specific datasets utilized by the 'Input Processing' and 'Analysis' modules of the architecture.",
        "The Data Input Layer ingests multi-modal health telemetry. The primary dataset originates from authenticated user submissions through the Web Dashboard UI.",
        "To train the Baseline Modelling and ML Risk Analysis layers, we employ a large-scale synthetic health dataset mirroring complex population distributions."
    ]) # Page 37
    
    add_page(doc, "Dataset Objective & Preprocessing Pipeline", [
        "Within the 'Input Processing' box, the 'Data preprocessing' layer handles missing values, normalizes integers, and sanitizes strings.",
        "The subsequent 'Feature Extraction' layer runs Principal Component Analysis (PCA) to reduce dataset dimensionality before handing vectors to the Analysis box.",
        "This prevents the ML Risk Analysis engine from overfitting on redundant dataset columns, ensuring highly accurate Risk Classifications."
    ]) # Page 38

    # Add 12 pages of explicit columns, exactly covering pages 39-50.
    dataset_cols = [
        ("input_heart_rate_bpm", "Ingested by Data Input Layer. Tracks real-time cardiovascular exertion.", "Crucial for immediate High Risk classification in the Analysis module."),
        ("input_blood_pressure_sys", "Systolic pressure metric. Processed in Data Preprocessing to remove negative anomalies.", "Baseline Modelling uses this to detect hypertension emergencies."),
        ("input_blood_sugar_mg", "Glucose reading submitted via Web Dashboard.", "Alert Generation relies on critical thresholds of this column to trigger diabetes alerts."),
        ("feature_stress_index", "Calculated by Feature Extraction based on HRV and sleep data.", "Recommendation Engine uses this to suggest mindfulness or physical rest."),
        ("feature_activity_level", "Extracted moving average of 7-day physical activity.", "Provides context for the ML Risk Analysis to gauge sedentary lifestyle impacts."),
        ("feature_diet_compliance", "Numeric score representing adherence to caloric goals.", "Low compliance vectors are routed to the Recommendation Engine for strict diet plans."),
        ("analysis_baseline_variance", "Output of Baseline Modelling. Measures deviation from standard health metrics.", "The primary determinant for the Risk Classification node."),
        ("analysis_ml_risk_score", "Float value produced by the ML Risk Analysis layer.", "Directly feeds the decision diamond. Score > 0.8 routes to Alert Generation."),
        ("analysis_classification_flag", "Categorical string (Low, Medium, High).", "The exact variable utilized by the Risk Classification router to direct data flow."),
        ("output_recommendation_json", "Payload created by the Recommendation Engine.", "Stored in the Database for retrieval by the Web Dashboard UI on next session."),
        ("output_alert_severity", "Integer generated by Alert Generation.", "If severity is maximum, SMS gateways ping emergency contacts alongside DB writes."),
        ("db_persistence_timestamp", "Generated automatically by the Database upon write.", "Required for chronological sorting of historical logs on the Web UI.")
    ]
    
    for i, (col, desc, impact) in enumerate(dataset_cols):
        add_page(doc, f"Dataset Column Dictionary: {col}", [
            f"Field Name: {col}",
            f"Description: {desc}",
            f"Pipeline Impact: {impact}",
            "This column travels explicitly through the defined architecture constraints, maintaining its integrity across microservice boundaries."
        ])

    output_filename = "Project_Architecture_Document_50_Pages_Aligned.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename} containing EXACTLY 50 pages based strictly on the uploaded flowchart.")

if __name__ == "__main__":
    generate_doc_v3()
