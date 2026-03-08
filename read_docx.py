import docx

doc_path = "Functional Document.docx"
try:
    with open("docx_output.txt", "w", encoding="utf-8") as f:
        doc = docx.Document(doc_path)
        for para in doc.paragraphs:
            f.write("P: " + para.text + "\n")
        f.write("--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for row in table.rows:
                row_data = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                f.write(str(row_data) + "\n")
            f.write("-" * 40 + "\n")
except Exception as e:
    with open("docx_output.txt", "w", encoding="utf-8") as f:
        f.write(f"Error reading docx: {e}")
